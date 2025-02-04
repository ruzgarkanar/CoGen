import pytesseract
from PIL import Image
from docx import Document
from pathlib import Path
from typing import Union, List, Optional
import PyPDF2
import io
import logging
from pdf2image import convert_from_path
import cv2
import numpy as np
import os

class DocumentExtractor:
    def __init__(self, ocr_lang: str = 'tur+eng', ocr_config: str = '--oem 3 --psm 3'):
        self.logger = logging.getLogger(__name__)
        self.ocr_lang = ocr_lang
        self.ocr_config = ocr_config
        self.temp_dir = Path("temp_ocr")
        self.temp_dir.mkdir(exist_ok=True)
        self.temp_files = []
        
        self.poppler_path = self._get_poppler_path()
        self.logger.info(f"Using poppler path: {self.poppler_path}")

    def _get_poppler_path(self) -> str:
        poppler_path = os.getenv('POPPLER_PATH')
        if (poppler_path and Path(poppler_path).exists()):
            return poppler_path
            
        if os.name == 'nt':  
            common_paths = [
                r"C:\Program Files\poppler-xx\bin",
                r"C:\Program Files (x86)\poppler-xx\bin",
                r"C:\poppler\bin"
            ]
        else:  
            common_paths = [
                "/opt/homebrew/bin",  
                "/usr/local/bin",     
                "/usr/bin",           
            ]
            
        for path in common_paths:
            if os.path.exists(path) and (
                os.path.exists(os.path.join(path, 'pdftoppm')) or 
                os.path.exists(os.path.join(path, 'pdftoppm.exe'))
            ):
                return path
                
        self.logger.warning("Poppler path not found in common locations")
        return None

    def extract_text(self, file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_from_pdf_with_fallback(file_path)
            elif file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_from_image_with_preprocessing(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            return ""

    def _extract_from_pdf_with_fallback(self, file_path: Path) -> str:
        text = ""
        
        try:
            convert_kwargs = {'dpi': 300}
            if self.poppler_path:
                convert_kwargs['poppler_path'] = self.poppler_path
                
            images = convert_from_path(file_path, **convert_kwargs)
            for page_num, image in enumerate(images):
                angle = self._detect_rotation(image)
                if (angle != 0):
                    image = image.rotate(angle)
                
                image = self._enhance_image_for_ocr(image)
                
                text += pytesseract.image_to_string(
                    image, 
                    lang=self.ocr_lang,
                    config=self.ocr_config
                )
        finally:
            self._cleanup_temp_files()

        return text.strip()

    def _preprocess_image(self, image: Image) -> Image:

        img_array = np.array(image)
        
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
            
        denoised = cv2.fastNlMeansDenoising(gray)
        
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        _, threshold = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
        final = cv2.morphologyEx(threshold, cv2.MORPH_OPEN, kernel)
        
        return Image.fromarray(final)

    def _extract_from_image_with_preprocessing(self, file_path: Path) -> str:
        try:
            img = cv2.imread(str(file_path))
            if img is None:
                raise ValueError("Failed to load image")

            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

            gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
            gray = cv2.dilate(gray, kernel, iterations=1)

            buf = io.BytesIO()
            cv2.imwrite("temp_processed.png", gray)
            
            text = pytesseract.image_to_string(Image.open("temp_processed.png"), lang=self.ocr_lang)
            if not text.strip():
                text = pytesseract.image_to_string(Image.open(file_path), lang=self.ocr_lang)
            
            return text

        except Exception as e:
            self.logger.error(f"Image OCR failed: {str(e)}")
            return ""

    def _extract_from_docx(self, file_path: Path) -> str:
        try:
            doc = Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {str(e)}")
            return ""

    def extract_from_directory(
        self, 
        directory_path: Union[str, Path],
        recursive: bool = True,
        file_types: Optional[List[str]] = None
    ) -> dict:

        if file_types is None:
            file_types = ['.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.docx']
            
        directory_path = Path(directory_path)
        results = {}
        
        self.logger.info(f"Scanning directory: {directory_path}")
        self.logger.info(f"Looking for file types: {file_types}")

        if not directory_path.exists():
            self.logger.error(f"Directory does not exist: {directory_path}")
            return results
            
        pattern = "**/*" if recursive else "*"
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in file_types:
                self.logger.info(f"Processing file: {file_path}")
                try:
                    text = self.extract_text(file_path)
                    if text and text.strip():
                        results[str(file_path)] = text
                        self.logger.info(f"Successfully extracted text from: {file_path}")
                    else:
                        self.logger.warning(f"No text extracted from: {file_path}")
                except Exception as e:
                    self.logger.error(f"Failed to process {file_path}: {str(e)}")
                    
        self.logger.info(f"Total files processed: {len(results)}")
        return results

    def _enhance_image_for_ocr(self, image: Image) -> Image:

        width, height = image.size
        image = image.resize((width*2, height*2), Image.LANCZOS)
        
        img_array = np.array(image)
        
        denoised = cv2.fastNlMeansDenoisingColored(img_array, None, 10, 10, 7, 21)
        
        kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
        sharpened = cv2.filter2D(denoised, -1, kernel)
        
        return Image.fromarray(sharpened)

    def _detect_rotation(self, image: Image) -> float:
        img_array = np.array(image)
        
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
            
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        
        coords = cv2.findNonZero(thresh)
        if coords is None:
            return 0.0
            
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = 90 + angle
        return -angle

    def _cleanup_temp_files(self):
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                print(f"Error cleaning up temporary file {temp_file}: {e}")
        self.temp_files = []
